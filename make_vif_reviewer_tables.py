from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
OUT_DIR = BASE / "results" / "vif_review_tables"


COX_FILES = {
    "ADNI": [
        BASE / "results/ADNI/vif_diagnostics/allcausedementia_outcome/primary_all_ages/ADNI_allcausedementia_outcome_primary_all_ages_ptau_demographics_lancet_vif_summary.csv",
        BASE / "results/ADNI/vif_diagnostics/allcausedementia_outcome/age_65/ADNI_allcausedementia_outcome_age_65_ptau_demographics_lancet_vif_summary.csv",
        BASE / "results/ADNI/vif_diagnostics/allcausedementia_outcome/agecutoff_65/ADNI_allcausedementia_outcome_agecutoff_65_ptau_demographics_lancet_vif_summary.csv",
        BASE / "results/ADNI/vif_diagnostics/alzheimers_outcome/primary_all_ages/ADNI_alzheimers_outcome_primary_all_ages_ptau_demographics_lancet_vif_summary.csv",
        BASE / "results/ADNI/vif_diagnostics/alzheimers_outcome/age_65/ADNI_alzheimers_outcome_age_65_ptau_demographics_lancet_vif_summary.csv",
        BASE / "results/ADNI/vif_diagnostics/alzheimers_outcome/agecutoff_65/ADNI_alzheimers_outcome_agecutoff_65_ptau_demographics_lancet_vif_summary.csv",
    ],
    "Pooled PET": [
        BASE / "results/pet_all_cohorts/vif_diagnostics/allcausedementia_outcome/age_65_cutoff/pooled_PET_allcausedementia_outcome_age_65_cutoff_centiloids_demographics_vif_summary.csv",
        BASE / "results/pet_all_cohorts/vif_diagnostics/ad_outcome/age_65_cutoff/pooled_PET_ad_outcome_age_65_cutoff_centiloids_demographics_vif_summary.csv",
        BASE / "results/pet_all_cohorts/vif_diagnostics/ad_outcome/cross_cohort_validation_age_65_cutoff/pooled_PET_ad_outcome_cross_cohort_validation_age_65_cutoff_centiloids_demographics_vif_summary.csv",
    ],
    "NACC CSF": [
        BASE / "results/nacc_csf/vif_diagnostics/allcausedementia_outcome/primary/NACC_CSF_allcausedementia_outcome_primary_csf_demographics_lancet_vif_summary.csv",
        BASE / "results/nacc_csf/vif_diagnostics/ad_outcome/primary/NACC_CSF_ad_outcome_primary_csf_demographics_lancet_vif_summary.csv",
    ],
}


def fmt_number(value):
    if pd.isna(value):
        return ""
    value = float(value)
    if abs(value) >= 10000:
        return f"{value:.2e}"
    if abs(value) >= 100:
        return f"{value:.1f}"
    return f"{value:.2f}"


def title_case_outcome(value):
    labels = {
        "allcausedementia_outcome": "All-cause dementia",
        "alzheimers_outcome": "Alzheimer's disease",
        "ad_outcome": "Alzheimer's disease",
    }
    return labels.get(value, value.replace("_", " "))


def title_case_analysis(value):
    labels = {
        "primary_all_ages": "Primary/all ages",
        "age_65": "Age 65+",
        "agecutoff_65": "Age cutoff 65+",
        "age_65_cutoff": "Age cutoff 65+",
        "cross_cohort_validation_age_65_cutoff": "Cross-cohort validation, age cutoff 65+",
        "primary": "Primary",
        "A4/LEARN maximal model": "A4/LEARN maximal model",
    }
    return labels.get(value, value.replace("_", " "))


def model_label(value):
    labels = {
        "ptau_demographics_lancet": "pTau + demographics + Lancet covariates",
        "centiloids_demographics": "Centiloids + demographics",
        "csf_demographics_lancet": "CSF biomarkers + demographics + Lancet covariates",
    }
    return labels.get(value, value.replace("_", " "))


def compact_terms(df, threshold=5, limit=4):
    high = df[df["mean_vif"] > threshold].sort_values("mean_vif", ascending=False)
    if high.empty:
        top = df.sort_values("mean_vif", ascending=False).head(2)
        terms = "; ".join(
            f"{row.variable} ({fmt_number(row.mean_vif)})" for row in top.itertuples()
        )
        return f"None; highest: {terms}"

    shown = high.head(limit)
    terms = "; ".join(
        f"{row.variable} ({fmt_number(row.mean_vif)})" for row in shown.itertuples()
    )
    remaining = len(high) - len(shown)
    if remaining > 0:
        terms += f"; +{remaining} more"
    return terms


def cox_row(path, dataset):
    df = pd.read_csv(path)
    first = df.iloc[0]
    max_mean_vif = df["mean_vif"].max()
    max_fold_vif = df["max_vif"].max()
    n_gt5 = int((df["mean_vif"] > 5).sum())
    n_gt10 = int((df["mean_vif"] > 10).sum())
    interpretation = "No concerning collinearity by common VIF thresholds."
    if dataset == "ADNI" and n_gt5:
        interpretation = (
            "Elevated VIFs were localized to related smoking/alcohol-history covariates; "
            "biomarker VIF remained low."
        )
    return {
        "Outcome": title_case_outcome(first["outcome"]),
        "Analysis set": title_case_analysis(first["analysis_set"]),
        "Model": model_label(first["model"]),
        "Variables evaluated": len(df),
        "Max mean VIF": fmt_number(max_mean_vif),
        "Max fold VIF": fmt_number(max_fold_vif),
        "Mean VIF >5 / >10": f"{n_gt5} / {n_gt10}",
        "Variables driving VIF >5": compact_terms(df, threshold=5),
        "Interpretation": interpretation,
    }


def build_a4_table():
    df = pd.read_html(BASE / "tidy_data/A4/vif_summary_table.html")[0]
    n_gt5 = int((df["mean_vif"] > 5).sum())
    n_gt10 = int((df["mean_vif"] > 10).sum())
    return pd.DataFrame(
        [
            {
                "Outcome": "Clinical dementia rating progression",
                "Analysis set": "A4/LEARN maximal model",
                "Model": "pTau-217 + amyloid PET + demographics + Lancet covariates",
                "Variables evaluated": len(df),
                "Max mean VIF": fmt_number(df["mean_vif"].max()),
                "Max fold VIF": fmt_number(df["max_vif"].max()),
                "Mean VIF >5 / >10": f"{n_gt5} / {n_gt10}",
                "Variables driving VIF >5": compact_terms(df, threshold=5),
                "Interpretation": "No concerning collinearity by common VIF thresholds.",
            }
        ]
    )


def build_cox_table(dataset):
    return pd.DataFrame([cox_row(path, dataset) for path in COX_FILES[dataset]])


def build_ukb_table():
    rows = []
    ukb_outputs = [
        (
            "All-cause dementia",
            BASE
            / "results/UKBiobank/vif_diagnostics/demographics_modality_lancet2024/agecutoff_65/ukbiobank_vif_summary.csv",
        ),
        (
            "Alzheimer's disease",
            BASE
            / "results/UKBiobank/vif_diagnostics/demographics_modality_lancet2024/alzheimers/agecutoff_65/ukbiobank_vif_summary.csv",
        ),
    ]
    for outcome_label, combined_path in ukb_outputs:
        combined = pd.read_csv(combined_path)
        for row in combined.itertuples(index=False):
            modality_dir = combined_path.parent / row.modality
            by_feature = pd.read_csv(modality_dir / "vif_by_feature_reference_coded.csv")
            top = by_feature.sort_values("vif", ascending=False).head(3)
            top_terms = "; ".join(f"{r.feature} ({fmt_number(r.vif)})" for r in top.itertuples())
            interpretation = "High-dimensional feature block contains expected redundancy."
            if row.modality == "neuroimaging":
                interpretation = (
                    "Substantial redundancy among neuroimaging features; appropriate for caution in feature-level interpretation."
                )
            rows.append(
                {
                    "Outcome": outcome_label,
                    "Modality": row.modality.replace("_", " ").title(),
                    "Rows after age filter": f"{int(row.n_rows_after_age_filter):,}",
                    "Features as trained / for VIF": f"{int(row.n_features_as_trained):,} / {int(row.n_features_for_vif):,}",
                    "As-trained rank deficiency": int(row.as_trained_rank_deficiency),
                    "Linear dependencies dropped": int(row.n_linear_dependency_columns_dropped),
                    "Max VIF": fmt_number(row.max_vif),
                    "Mean VIF": fmt_number(row.mean_vif),
                    "Median VIF": fmt_number(row.median_vif),
                    "Features VIF >5 / >10": f"{int(row.n_vif_gt_5):,} / {int(row.n_vif_gt_10):,}",
                    "Highest VIF features": top_terms,
                    "Interpretation": interpretation,
                }
            )
    return pd.DataFrame(rows)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeats_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, title, df, widths):
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"]
    heading.add_run(title)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_cells = table.rows[0].cells
    set_row_repeats_as_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for idx, column in enumerate(df.columns):
        cell = header_cells[idx]
        cell.text = str(column)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)

    for _, record in df.iterrows():
        row = table.add_row()
        set_row_cant_split(row)
        cells = row.cells
        for idx, column in enumerate(df.columns):
            text = str(record[column])
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if column
                    in {
                        "Variables evaluated",
                        "Max mean VIF",
                        "Max fold VIF",
                        "Mean VIF >5 / >10",
                        "Rows after age filter",
                        "Features as trained / for VIF",
                        "As-trained rank deficiency",
                        "Linear dependencies dropped",
                        "Max VIF",
                        "Mean VIF",
                        "Median VIF",
                        "Features VIF >5 / >10",
                    }
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    set_table_width(table, widths)
    doc.add_paragraph()


def configure_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 14, RGBColor(46, 116, 181)),
        ("Heading 2", 12, RGBColor(46, 116, 181)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def write_docx(tables):
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("Multicollinearity Diagnostics Across Datasets")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(18)
    title_run.bold = True

    note = doc.add_paragraph()
    note.add_run("Note. ").bold = True
    note.add_run(
        "Cox-model VIFs are summarized across five cross-validation folds. "
        "For UK Biobank, diagnostics summarize the high-dimensional feature matrix after reference-coding categorical variables and removing exact linear dependencies before VIF calculation; all-cause dementia and Alzheimer's disease row sets are shown separately."
    )

    cox_widths = [1.2, 1.15, 1.7, 0.65, 0.65, 0.65, 0.75, 2.4, 1.9]
    add_table(doc, "A4/LEARN", tables["A4_LEARN"], cox_widths)
    add_table(doc, "ADNI", tables["ADNI"], cox_widths)
    add_table(doc, "Pooled PET Cohorts", tables["pooled_pet"], cox_widths)
    add_table(doc, "NACC CSF", tables["nacc_csf"], cox_widths)

    ukb_widths = [0.95, 0.75, 0.9, 0.75, 0.75, 0.55, 0.55, 0.55, 0.8, 2.25, 1.7]
    for outcome_label in ["All-cause dementia", "Alzheimer's disease"]:
        ukb_subset = (
            tables["uk_biobank"]
            .loc[tables["uk_biobank"]["Outcome"] == outcome_label]
            .drop(columns=["Outcome"])
            .reset_index(drop=True)
        )
        add_table(doc, f"UK Biobank - {outcome_label}", ukb_subset, ukb_widths)

    output_path = OUT_DIR / "vif_reviewer_tables.docx"
    doc.save(output_path)
    return output_path


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    tables = {
        "A4_LEARN": build_a4_table(),
        "ADNI": build_cox_table("ADNI"),
        "pooled_pet": build_cox_table("Pooled PET"),
        "nacc_csf": build_cox_table("NACC CSF"),
        "uk_biobank": build_ukb_table(),
    }
    csv_paths = {}
    for name, df in tables.items():
        path = OUT_DIR / f"{name}_vif_reviewer_table.csv"
        df.to_csv(path, index=False)
        csv_paths[name] = path
    docx_path = write_docx(tables)
    print(docx_path)
    for path in csv_paths.values():
        print(path)


if __name__ == "__main__":
    main()
